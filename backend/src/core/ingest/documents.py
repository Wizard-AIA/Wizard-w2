"""Reference documents attached to a session.

A data dictionary, a fee schedule, a metric definition, the page explaining that
``status = 'C'`` means cancelled and not complete. Hard analytical questions turn
on these far more often than on anything discoverable from the tables, and until
now there was nowhere to put one: ingestion accepted tabular files only.

`DABstep <https://arxiv.org/abs/2506.23719>`_ is built around exactly this —
its hard tasks require cross-referencing structured data against unstructured
documentation, and one of its named failure modes is agents failing "to consult
required documentation at the right point".

Design
------
Deliberately small. Documents are chunked on paragraph boundaries, embedded once
at upload, and retrieved by the same :mod:`~src.core.embeddings` service the rest
of the app uses -- which degrades to lexical overlap when no transformer is
loaded, so this works air-gapped and in CI.

PDF and DOCX parsing is *optional*. The extractors are imported inside the
functions that need them, so a deployment that never uploads a PDF does not pay
for the dependency and an install without it still starts.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from src.config import settings
from src.core.embeddings import embedding_service
from src.utils.logging import logger


TEXT_LIKE = {".md", ".markdown", ".txt", ".rst", ".text", ".log"}
PDF_LIKE = {".pdf"}
DOCX_LIKE = {".docx"}
HTML_LIKE = {".html", ".htm"}

SUPPORTED_DOCUMENT_EXTENSIONS = TEXT_LIKE | PDF_LIKE | DOCX_LIKE | HTML_LIKE

#: Paragraph boundary: one or more blank lines, or a markdown heading.
PARAGRAPH_BREAK = re.compile(r"\n\s*\n+")
HTML_TAG = re.compile(r"<[^>]+>")
WHITESPACE = re.compile(r"[ \t]+")


class UnsupportedDocumentError(ValueError):
    """Raised for a document extension that cannot be parsed."""


class DocumentExtractionError(RuntimeError):
    """Raised when a supported format is present but its parser is not installed."""


@dataclass
class DocumentChunk:
    """One retrievable passage."""

    document: str
    index: int
    text: str
    embedding: list[float] | None = None

    def to_dict(self) -> dict[str, Any]:
        return {"document": self.document, "index": self.index, "text": self.text}


@dataclass
class ContextDocument:
    """A reference document belonging to one session."""

    name: str
    text: str
    chunks: list[DocumentChunk] = field(default_factory=list)
    source_format: str = "txt"

    @property
    def char_count(self) -> int:
        return len(self.text)

    def summary(self) -> dict[str, Any]:
        head = self.text.strip().splitlines()
        return {
            "name": self.name,
            "chars": self.char_count,
            "chunks": len(self.chunks),
            "source_format": self.source_format,
            "preview": (head[0][:200] if head else ""),
        }


# ---------------------------------------------------------------------- #
# Extraction
# ---------------------------------------------------------------------- #
def _extract_text(path: Path, suffix: str) -> str:
    if suffix in TEXT_LIKE:
        return _read_text(path)

    if suffix in HTML_LIKE:
        return HTML_TAG.sub(" ", _read_text(path))

    if suffix in PDF_LIKE:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - depends on optional extra
            raise DocumentExtractionError(
                "Reading PDFs needs the `pypdf` package. Install it, or upload the document as Markdown or text."
            ) from exc
        reader = PdfReader(str(path))

        # A small file on disk does not bound what comes out of it: a PDF can
        # declare an arbitrary page count, and a page's content stream can
        # decompress to far more text than its own size suggests. Both are
        # cheap for an attacker and expensive for this process, so both are
        # capped before extraction is allowed to run unbounded.
        page_count = len(reader.pages)
        if page_count > settings.CONTEXT_DOC_MAX_PDF_PAGES:
            raise DocumentExtractionError(
                f"'{path.name}' has {page_count} pages, more than the "
                f"{settings.CONTEXT_DOC_MAX_PDF_PAGES} this deployment will parse."
            )

        parts: list[str] = []
        total_chars = 0
        for page in reader.pages:
            text = page.extract_text() or ""
            total_chars += len(text)
            if total_chars > settings.CONTEXT_DOC_MAX_EXTRACTED_CHARS:
                raise DocumentExtractionError(
                    f"'{path.name}' decompresses to more text than this deployment will hold in memory "
                    f"(over {settings.CONTEXT_DOC_MAX_EXTRACTED_CHARS:,} characters)."
                )
            parts.append(text)
        return "\n\n".join(parts)

    if suffix in DOCX_LIKE:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - depends on optional extra
            raise DocumentExtractionError(
                "Reading .docx needs the `python-docx` package. Install it, or upload the document as Markdown or text."
            ) from exc
        document = docx.Document(str(path))
        return "\n\n".join(paragraph.text for paragraph in document.paragraphs)

    raise UnsupportedDocumentError(f"Unsupported document type '{suffix}'.")


def _read_text(path: Path) -> str:
    """Decodes a text file, tolerating whatever encoding it arrived in.

    A data dictionary exported from Excel on Windows is cp1252 far more often
    than it is UTF-8, and failing the upload over a smart quote would be absurd.
    """
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------- #
# Chunking
# ---------------------------------------------------------------------- #
def chunk_text(text: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    """Splits on paragraph boundaries, packing up to ``size`` characters.

    Paragraph-aligned rather than fixed-width because these documents are
    definitions and rules: cutting one in half produces two chunks that each
    retrieve well and neither of which states the rule.
    """
    limit = size or settings.CONTEXT_CHUNK_CHARS
    lap = overlap if overlap is not None else settings.CONTEXT_CHUNK_OVERLAP

    normalised = WHITESPACE.sub(" ", (text or "").replace("\r\n", "\n")).strip()
    if not normalised:
        return []

    paragraphs = [p.strip() for p in PARAGRAPH_BREAK.split(normalised) if p.strip()]
    if not paragraphs:
        return []

    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        # A paragraph longer than the budget is split on its own, by sentence.
        if len(paragraph) > limit:
            if current:
                chunks.append(current)
                current = ""
            chunks.extend(_split_long(paragraph, limit))
            continue

        candidate = f"{current}\n\n{paragraph}" if current else paragraph
        if len(candidate) <= limit:
            current = candidate
        else:
            chunks.append(current)
            # Carry the tail of the previous chunk so a rule spanning the
            # boundary is still retrievable from either side.
            current = (current[-lap:] + "\n\n" + paragraph) if lap else paragraph
    if current:
        chunks.append(current)
    return chunks


def _split_long(paragraph: str, limit: int) -> list[str]:
    sentences = re.split(r"(?<=[.!?])\s+", paragraph)
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        if len(sentence) > limit:
            # A single sentence over the budget: hard-cut it, nothing else to do.
            for start in range(0, len(sentence), limit):
                chunks.append(sentence[start : start + limit])
            continue
        candidate = f"{current} {sentence}".strip()
        if len(candidate) <= limit:
            current = candidate
        else:
            if current:
                chunks.append(current)
            current = sentence
    if current:
        chunks.append(current)
    return chunks


# ---------------------------------------------------------------------- #
# Loading
# ---------------------------------------------------------------------- #
def supported_document_extensions() -> list[str]:
    return sorted(SUPPORTED_DOCUMENT_EXTENSIONS)


def is_supported_document(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS


def load_document(path: Path, name: str) -> ContextDocument:
    """Parses, chunks and embeds one reference document."""
    suffix = Path(name).suffix.lower()
    if suffix not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise UnsupportedDocumentError(
            f"Unsupported document type '{suffix or name}'. Supported: {', '.join(supported_document_extensions())}"
        )

    text = _extract_text(path, suffix)
    if not text.strip():
        raise UnsupportedDocumentError(f"'{name}' parsed successfully but contains no readable text.")

    document = ContextDocument(name=name, text=text, source_format=suffix.lstrip("."))
    for index, body in enumerate(chunk_text(text)):
        embedding = None
        try:
            embedding = embedding_service.encode(body)
        except Exception as exc:  # pragma: no cover - encoder degrades on its own
            logger.debug("Chunk embedding failed; lexical fallback will be used", error=str(exc))
        document.chunks.append(DocumentChunk(document=name, index=index, text=body, embedding=embedding))

    logger.info("Context document loaded", document=name, chars=len(text), chunks=len(document.chunks))
    return document


def search_documents(
    documents: dict[str, ContextDocument],
    query: str,
    limit: int | None = None,
) -> list[tuple[str, str]]:
    """Ranks every chunk across every document. Returns ``(document, passage)``.

    Ranking goes through :mod:`~src.core.embeddings`, which uses a transformer
    when one is loaded and falls back to lexical overlap when none is -- so
    retrieval works in an air-gapped install, just less well.
    """
    top_k = limit or settings.CONTEXT_TOP_K
    chunks = [chunk for document in documents.values() for chunk in document.chunks]
    if not chunks or not query.strip():
        return []

    ranked = embedding_service.rank(query, [(chunk.text, chunk.embedding) for chunk in chunks])

    results: list[tuple[str, str]] = []
    for score, index in ranked[:top_k]:
        if score < settings.RAG_MIN_SIMILARITY:
            continue
        chunk = chunks[index]
        results.append((chunk.document, chunk.text))
    return results
